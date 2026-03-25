from __future__ import annotations
import os
from typing import Optional

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

SUPPORTED_EXTS = {".pdf", ".docx", ".txt"}

def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)

def parse_resume_file(uploaded_file) -> str:
    """Parse uploaded resume file (PDF/DOCX/TXT) to raw text.
    - Works with InMemoryUploadedFile / TemporaryUploadedFile.
    """
    name = getattr(uploaded_file, "name", "resume")
    ext = os.path.splitext(name)[1].lower()

    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {sorted(SUPPORTED_EXTS)}")

    # Save to temp path for libraries that need a filename
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        for chunk in uploaded_file.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name

    try:
        if ext == ".pdf":
            text = pdf_extract_text(tmp_path) or ""
        elif ext == ".docx":
            text = _read_docx(tmp_path) or ""
        else:
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        return text.strip()
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
